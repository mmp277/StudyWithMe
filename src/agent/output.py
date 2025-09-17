from typing import List, Tuple, Dict
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def write_summary_docx(output_path: str, file_summaries: List[Tuple[str, List[str]]]) -> None:
	doc = Document()
	doc.add_heading("Lecture Summaries", level=1)
	for filename, summaries in file_summaries:
		title = os.path.basename(filename)
		doc.add_heading(title, level=2)
		for s in summaries:
			p = doc.add_paragraph(s)
			p_format = p.paragraph_format
			p_format.space_after = Pt(6)
		doc.add_page_break()
	doc.save(output_path)


def write_flashcards_docx(output_path: str, flashcards_by_file: List[Tuple[str, List[Tuple[str, str]]]]) -> None:
	doc = Document()
	doc.add_heading("Flashcards", level=1)
	for filename, cards in flashcards_by_file:
		title = os.path.basename(filename)
		doc.add_heading(title, level=2)
		for idx, (q, a) in enumerate(cards, start=1):
			q_p = doc.add_paragraph()
			run_q = q_p.add_run(f"Q{idx}. {q}")
			run_q.bold = True
			a_p = doc.add_paragraph(f"A{idx}. {a}")
			a_p_format = a_p.paragraph_format
			a_p_format.space_after = Pt(6)
		doc.add_page_break()
	doc.save(output_path)


def write_formula_sheet_docx(output_path: str, formulas_by_file: List[Tuple[str, List[Tuple[str, Dict[str, str]]]]]) -> None:
	doc = Document()
	doc.add_heading("Formula Sheet", level=1)
	for filename, items in formulas_by_file:
		title = os.path.basename(filename)
		doc.add_heading(title, level=2)
		for eq, defs in items:
			p_eq = doc.add_paragraph()
			r = p_eq.add_run(eq)
			r.italic = True
			p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

			if defs:
				doc.add_paragraph("Variables:")
				for sym, desc in defs.items():
					doc.add_paragraph(f"{sym}: {desc}", style=None)
			doc.add_paragraph("")
		doc.add_page_break()
	doc.save(output_path)
